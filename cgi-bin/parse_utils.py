#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件解析工具类
支持: .txt, .csv, .xlsx, .xls, .doc, .docx
"""

import os
import re
import struct

class ParseUtils:
    """文件解析工具类"""

    def __init__(self, file_path, ext):
        self.file_path = file_path
        self.ext = ext.lower()

    def extract_numbers(self):
        """根据文件类型调用相应的解析方法"""
        if self.ext == '.txt':
            return self._parse_txt()
        elif self.ext == '.csv':
            return self._parse_csv()
        elif self.ext == '.xlsx':
            return self._parse_xlsx()
        elif self.ext == '.xls':
            return self._parse_xls()
        elif self.ext == '.doc':
            return self._parse_doc()
        elif self.ext == '.docx':
            return self._parse_docx()
        else:
            raise ValueError(f"不支持的文件格式: {self.ext}")

    def _parse_txt(self):
        """解析纯文本文件"""
        numbers = []
        with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                numbers.extend(self._extract_numbers_from_text(line))
        return numbers

    def _parse_csv(self):
        """解析 CSV 文件"""
        numbers = []
        with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                # 分割 CSV 行（简单实现，复杂情况需用 csv 模块）
                parts = line.strip().split(',')
                for part in parts:
                    numbers.extend(self._extract_numbers_from_text(part))
        return numbers

    def _parse_xlsx(self):
        """解析 XLSX 文件 (Excel 2007+)"""
        try:
            import openpyxl
        except ImportError:
            raise ImportError("需要安装 openpyxl: pip install openpyxl")

        numbers = []
        wb = openpyxl.load_workbook(self.file_path, data_only=True)

        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        numbers.extend(self._extract_numbers_from_text(str(cell.value)))

        return numbers

    def _parse_xls(self):
        """解析 XLS 文件 (Excel 97-2003)"""
        try:
            import xlrd
        except ImportError:
            raise ImportError("需要安装 xlrd: pip install xlrd")

        numbers = []
        wb = xlrd.open_workbook(self.file_path)

        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    cell_value = sheet.cell_value(row_idx, col_idx)
                    if cell_value is not None:
                        numbers.extend(self._extract_numbers_from_text(str(cell_value)))

        return numbers

    def _parse_doc(self):
        """解析 DOC 文件 (Word 97-2003)"""
        try:
            import win32com.client
        except ImportError:
            # 尝试使用 antiword
            import subprocess
            try:
                result = subprocess.run(['antiword', self.file_path],
                                       capture_output=True, text=True)
                if result.returncode == 0:
                    text = result.stdout
                    return self._extract_numbers_from_text(text)
            except FileNotFoundError:
                pass

            raise ImportError("需要安装 pywin32 或 antiword 来解析 .doc 文件")

        # 使用 Word COM 自动化
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            doc = word.Documents.Open(os.path.abspath(self.file_path))
            text = doc.Content.Text
            doc.Close()
            return self._extract_numbers_from_text(text)
        finally:
            word.Quit()

    def _parse_docx(self):
        """解析 DOCX 文件 (Word 2007+)"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        numbers = []
        doc = Document(self.file_path)

        # 提取段落文本
        for para in doc.paragraphs:
            numbers.extend(self._extract_numbers_from_text(para.text))

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    numbers.extend(self._extract_numbers_from_text(cell.text))

        return numbers

    def _extract_numbers_from_text(self, text):
        """从文本中提取所有数字（整数和浮点数）"""
        if not text:
            return []

        # 匹配数字，包括负数、浮点数
        # 匹配模式: 可选的负号 + 数字部分 + 可选的小数部分
        pattern = r'-?\d+\.?\d*'

        numbers = []
        for match in re.finditer(pattern, text):
            num_str = match.group()
            try:
                num = float(num_str)
                # 如果是整数，转换为 int
                if num == int(num):
                    numbers.append(int(num))
                else:
                    numbers.append(num)
            except ValueError:
                continue

        return numbers
